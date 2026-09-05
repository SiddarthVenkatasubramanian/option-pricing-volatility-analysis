"""TCS Options Pricing & Volatility Analysis

CRR Binomial and Black-Scholes pricing using TCS market data, with historical
and implied volatility, liquidity-aware strike selection, Greeks, sensitivity,
put-call parity, convergence analysis, run-history logging, charts, and
automatic Word-report generation.

Run:
    python tcs_options_pricing.py --help
"""

import argparse
import csv
import os
from datetime import date, datetime, timedelta

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from scipy.stats import norm

try:
    import yfinance as yf
    HAVE_YFINANCE = True
except ImportError:
    HAVE_YFINANCE = False

try:
    from docx import Document
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False


SCRIPT_VERSION = "1.5.0"
EXERCISE_STYLE = "European"


# =========================================================
# 1. EXPIRY
# =========================================================

NSE_HOLIDAYS_2026 = {
    date(2026, 1, 26), date(2026, 3, 3), date(2026, 3, 26),
    date(2026, 3, 31), date(2026, 4, 3), date(2026, 4, 14),
    date(2026, 5, 1), date(2026, 5, 28), date(2026, 6, 26),
    date(2026, 9, 14), date(2026, 10, 2), date(2026, 10, 20),
    date(2026, 11, 10), date(2026, 11, 24), date(2026, 12, 25),
}
NSE_HOLIDAY_CALENDARS = {2026: NSE_HOLIDAYS_2026}


def is_nse_trading_holiday(d):
    return d in NSE_HOLIDAY_CALENDARS.get(d.year, frozenset())


def roll_back_to_trading_day(d):
    """Roll a scheduled expiry backward if it is a weekend or known holiday."""
    if d.year not in NSE_HOLIDAY_CALENDARS:
        print(
            f"[WARN] No NSE holiday calendar loaded for {d.year}; "
            "holiday adjustment skipped. Pass --expiry explicitly."
        )
        return d

    while d.weekday() >= 5 or is_nse_trading_holiday(d):
        d -= timedelta(days=1)
    return d


def compute_monthly_expiry(run_date, expiry_weekday=1):
    """Return the last configured weekday of the current/following month."""
    def last_weekday(year, month, weekday):
        if month == 12:
            first_next = date(year + 1, 1, 1)
        else:
            first_next = date(year, month + 1, 1)
        last_day = first_next - timedelta(days=1)
        return last_day - timedelta(days=(last_day.weekday() - weekday) % 7)

    year, month = run_date.year, run_date.month
    expiry = roll_back_to_trading_day(last_weekday(year, month, expiry_weekday))

    if expiry <= run_date:
        month += 1
        if month == 13:
            year += 1
            month = 1
        expiry = roll_back_to_trading_day(last_weekday(year, month, expiry_weekday))

    return expiry


def resolve_expiry(run_date, override, expiry_weekday):
    if override:
        expiry = date.fromisoformat(override)
    else:
        expiry = compute_monthly_expiry(run_date, expiry_weekday)
        status = (
            "NSE-holiday-adjusted"
            if expiry.year in NSE_HOLIDAY_CALENDARS
            else "not holiday-adjusted"
        )
        print(
            f"[INFO] --expiry not supplied; auto-computed {expiry} ({status}). "
            "Pass --expiry explicitly if it does not match the downloaded chain."
        )

    if expiry <= run_date:
        raise ValueError(f"Expiry {expiry} is not after run date {run_date}.")
    return expiry


# =========================================================
# 2. HISTORICAL VOLATILITY
# =========================================================

def get_historical_volatility(symbol, end_date=None, lookback_days=365):
    """Download one calendar year's prices ending at end_date."""
    if not HAVE_YFINANCE:
        raise RuntimeError("yfinance is not installed.")

    end_date = end_date or date.today()
    start_date = end_date - timedelta(days=lookback_days)
    download_start = start_date - timedelta(days=7)

    data = yf.download(
        symbol,
        start=download_start.isoformat(),
        end=(end_date + timedelta(days=1)).isoformat(),
        auto_adjust=False,
        progress=False,
    )["Close"].squeeze()

    data = data.dropna()
    if data.empty:
        raise ValueError(f"No valid price data downloaded for {symbol}.")

    returns = np.log(data / data.shift(1)).dropna()
    returns = returns[returns.index.date >= start_date]
    returns = returns[returns.index.date <= end_date]

    if returns.empty:
        raise ValueError("No returns available for the requested volatility window.")

    mean_return = float(returns.mean() * 252)
    sigma = float(returns.std() * np.sqrt(252))
    data_start = returns.index.min().date()
    return returns, mean_return, sigma, data_start


# =========================================================
# 3. CRR BINOMIAL MODEL
# =========================================================

def binomial_price(S, K, r, sigma, T, n, option_type="call", q=0.0):
    """European CRR binomial option price with continuous dividend yield q."""
    if EXERCISE_STYLE != "European":
        raise ValueError("This implementation supports European exercise only.")
    if T <= 0 or sigma <= 0 or n < 1:
        raise ValueError("Require T>0, sigma>0 and n>=1.")

    dt = T / n
    u = np.exp(sigma * np.sqrt(dt))
    d = 1.0 / u
    p = (np.exp((r - q) * dt) - d) / (u - d)

    if not 0 < p < 1:
        raise ValueError("Risk-neutral probability p is outside (0,1).")

    j = np.arange(n + 1)
    ST = S * (u ** j) * (d ** (n - j))

    if option_type == "call":
        values = np.maximum(ST - K, 0.0)
    elif option_type == "put":
        values = np.maximum(K - ST, 0.0)
    else:
        raise ValueError("option_type must be 'call' or 'put'.")

    discount = np.exp(-r * dt)
    for i in range(n - 1, -1, -1):
        values = discount * (
            p * values[1:i + 2] + (1 - p) * values[0:i + 1]
        )

    return float(values[0])


# =========================================================
# 4. BLACK-SCHOLES + GREEKS
# =========================================================

def black_scholes(S, K, r, sigma, T, option_type="call", q=0.0):
    """European Black-Scholes price with continuous dividend yield q."""
    if T <= 0 or sigma <= 0:
        raise ValueError("Require T>0 and sigma>0.")

    d1 = (
        np.log(S / K) + (r - q + 0.5 * sigma ** 2) * T
    ) / (sigma * np.sqrt(T))
    d2 = d1 - sigma * np.sqrt(T)

    call = (
        S * np.exp(-q * T) * norm.cdf(d1)
        - K * np.exp(-r * T) * norm.cdf(d2)
    )
    put = (
        K * np.exp(-r * T) * norm.cdf(-d2)
        - S * np.exp(-q * T) * norm.cdf(-d1)
    )

    if option_type == "call":
        return float(call)
    if option_type == "put":
        return float(put)
    raise ValueError("option_type must be 'call' or 'put'.")


def bs_greeks(S, K, r, sigma, T, q=0.0):
    """BS Greeks: vega per 1% vol move, theta per day, rho per 1% rate move."""
    if T <= 0 or sigma <= 0:
        raise ValueError("Require T>0 and sigma>0.")

    d1 = (
        np.log(S / K) + (r - q + 0.5 * sigma ** 2) * T
    ) / (sigma * np.sqrt(T))
    d2 = d1 - sigma * np.sqrt(T)

    return {
        "delta_call": float(np.exp(-q * T) * norm.cdf(d1)),
        "delta_put": float(np.exp(-q * T) * (norm.cdf(d1) - 1)),
        "gamma": float(np.exp(-q * T) * norm.pdf(d1) / (S * sigma * np.sqrt(T))),
        "vega": float(S * np.exp(-q * T) * norm.pdf(d1) * np.sqrt(T) / 100),
        "theta_call": float((
            -S * np.exp(-q * T) * norm.pdf(d1) * sigma / (2 * np.sqrt(T))
            - r * K * np.exp(-r * T) * norm.cdf(d2)
            + q * S * np.exp(-q * T) * norm.cdf(d1)
        ) / 365),
        "theta_put": float((
            -S * np.exp(-q * T) * norm.pdf(d1) * sigma / (2 * np.sqrt(T))
            + r * K * np.exp(-r * T) * norm.cdf(-d2)
            - q * S * np.exp(-q * T) * norm.cdf(-d1)
        ) / 365),
        "rho_call": float(K * T * np.exp(-r * T) * norm.cdf(d2) / 100),
        "rho_put": float(-K * T * np.exp(-r * T) * norm.cdf(-d2) / 100),
    }


# =========================================================
# 5. IMPLIED VOLATILITY
# =========================================================

def implied_volatility(
    market_price, S, K, r, T, option_type="call", q=0.0,
    tol=1e-6, max_iter=100, sigma_bounds=(1e-4, 5.0),
):
    """Solve IV with Newton-Raphson and bisection fallback."""
    if not np.isfinite(market_price) or market_price <= 0:
        return np.nan

    lo, hi = sigma_bounds
    price_lo = black_scholes(S, K, r, lo, T, option_type, q)
    price_hi = black_scholes(S, K, r, hi, T, option_type, q)
    if not price_lo - tol <= market_price <= price_hi + tol:
        return np.nan

    sigma = 0.30
    for _ in range(max_iter):
        price = black_scholes(S, K, r, sigma, T, option_type, q)
        vega = bs_greeks(S, K, r, sigma, T, q)["vega"] * 100
        diff = price - market_price

        if abs(diff) < tol:
            return float(sigma)
        if vega < 1e-8:
            break

        sigma -= diff / vega
        if not lo < sigma < hi:
            break

    # Guaranteed monotonic fallback for standard European call/put prices.
    a, b = lo, hi
    for _ in range(200):
        mid = (a + b) / 2
        price_mid = black_scholes(S, K, r, mid, T, option_type, q)
        if abs(price_mid - market_price) < tol:
            return float(mid)
        if price_mid < market_price:
            a = mid
        else:
            b = mid

    return float((a + b) / 2)


# =========================================================
# 6. OPTION CHAIN + STRIKE SELECTION
# =========================================================

def load_option_chain(path):
    """Load a standard NSE option-chain CSV and normalize its key columns."""
    df = pd.read_csv(path, skiprows=1, encoding="utf-8-sig")
    df.columns = df.columns.str.strip()

    required = {
        "STRIKE": "strike",
        "LTP": "call_ltp",
        "IV": "call_iv",
        "LTP.1": "put_ltp",
        "IV.1": "put_iv",
    }
    optional = {
        "OI": "call_oi",
        "VOLUME": "call_volume",
        "OI.1": "put_oi",
        "VOLUME.1": "put_volume",
    }

    missing = [c for c in required if c not in df.columns]
    if missing:
        raise ValueError(
            f"Missing required option-chain columns: {missing}. "
            f"Columns found: {list(df.columns)}"
        )

    columns = {**required, **{k: v for k, v in optional.items() if k in df.columns}}
    chain = df[list(columns)].rename(columns=columns).copy()

    for col in chain.columns:
        chain[col] = pd.to_numeric(
            chain[col].astype(str)
            .str.replace(",", "")
            .str.replace("%", "")
            .str.strip()
            .str.replace(r"^-+$", "", regex=True),
            errors="coerce",
        )

    chain = chain.dropna(subset=["strike"]).sort_values("strike")
    chain["call_iv"] /= 100
    chain["put_iv"] /= 100

    liquidity_info = {
        "call": {"call_oi", "call_volume"}.issubset(chain.columns),
        "put": {"put_oi", "put_volume"}.issubset(chain.columns),
    }
    return chain, liquidity_info


def select_atm_strike(chain, S, have_liquidity_cols, min_oi=100, min_volume=1):
    """Select the nearest usable call strike, preferring liquid quotes."""
    valid = chain.dropna(subset=["call_ltp", "call_iv"]).copy()
    valid = valid[(valid["call_ltp"] > 0) & (valid["call_iv"] > 0.01)]

    if have_liquidity_cols:
        liquid = valid[
            (valid["call_oi"].fillna(0) >= min_oi)
            & (valid["call_volume"].fillna(0) >= min_volume)
        ]
        if not liquid.empty:
            row = liquid.loc[(liquid["strike"] - S).abs().idxmin()]
            return row, True

    if valid.empty:
        raise ValueError("No usable call LTP/IV quotes available for strike selection.")

    # Fallback: nearest quote with a reasonable positive IV.
    row = valid.loc[(valid["strike"] - S).abs().idxmin()]
    return row, False


def filter_trustworthy_strikes(
    chain, iv_col, oi_col, volume_col, has_liquidity,
    reference_strike, min_oi=100, min_volume=1,
    moneyness_bound=0.15, iv_ceiling=1.50,
):
    side = chain.dropna(subset=[iv_col]).copy()
    if has_liquidity:
        side = side[
            (side[oi_col].fillna(0) >= min_oi)
            & (side[volume_col].fillna(0) >= min_volume)
        ]
    return side[
        (side[iv_col] > 0.01)
        & (side[iv_col] <= iv_ceiling)
        & ((side["strike"] - reference_strike).abs() <= moneyness_bound * reference_strike)
    ]


def compute_skew_stats(chain, liquidity_info, atm_strike, min_oi=100, min_volume=1):
    calls = filter_trustworthy_strikes(
        chain, "call_iv", "call_oi", "call_volume", liquidity_info["call"],
        atm_strike, min_oi, min_volume,
    )
    puts = filter_trustworthy_strikes(
        chain, "put_iv", "put_oi", "put_volume", liquidity_info["put"],
        atm_strike, min_oi, min_volume,
    )

    both = chain.dropna(subset=["call_iv", "put_iv"]).copy()
    if liquidity_info["call"] and liquidity_info["put"]:
        both = both[
            (both["call_oi"].fillna(0) >= min_oi)
            & (both["call_volume"].fillna(0) >= min_volume)
            & (both["put_oi"].fillna(0) >= min_oi)
            & (both["put_volume"].fillna(0) >= min_volume)
        ]
    both = both[
        (both["call_iv"] > 0.01) & (both["call_iv"] <= 1.50)
        & (both["put_iv"] > 0.01) & (both["put_iv"] <= 1.50)
        & ((both["strike"] - atm_strike).abs() <= 0.15 * atm_strike)
    ]

    stats = {
        "call_iv_min": calls["call_iv"].min() if not calls.empty else None,
        "call_iv_max": calls["call_iv"].max() if not calls.empty else None,
        "put_iv_min": puts["put_iv"].min() if not puts.empty else None,
        "put_iv_max": puts["put_iv"].max() if not puts.empty else None,
        "filtered": liquidity_info["call"] and liquidity_info["put"],
        "call_filtered": liquidity_info["call"],
        "put_filtered": liquidity_info["put"],
    }

    if not both.empty:
        row = both.loc[(both["strike"] - atm_strike).abs().idxmin()]
        gap = float(row["put_iv"] - row["call_iv"])
        stats["atm_gap"] = gap
        stats["atm_gap_strike"] = float(row["strike"])
    else:
        stats["atm_gap"] = None
        stats["atm_gap_strike"] = None

    return stats


# =========================================================
# 7. RUN HISTORY
# =========================================================

def log_run(log_path, record):
    """Append a run, migrating older schemas when necessary."""
    fields = list(record.keys())

    if not os.path.isfile(log_path):
        with open(log_path, "w", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=fields)
            writer.writeheader()
            writer.writerow(record)
        return

    with open(log_path, "r", newline="") as f:
        rows = list(csv.reader(f))

    if not rows:
        header, old_rows = [], []
    else:
        header, old_rows = rows[0], rows[1:]

    if header == fields and all(len(row) == len(fields) for row in old_rows):
        with open(log_path, "a", newline="") as f:
            csv.DictWriter(f, fieldnames=fields).writerow(record)
        return

    unified = fields + [c for c in header if c not in fields]
    migrated = []
    for row in old_rows:
        if len(row) == len(header):
            migrated.append(dict(zip(header, row)))
        elif len(row) == len(fields):
            migrated.append(dict(zip(fields, row)))
        else:
            migrated.append({f"_unrecognised_col_{i}": value for i, value in enumerate(row)})

    print(f"[WARN] {log_path} schema changed or was inconsistent; migrating {len(migrated)} row(s).")

    with open(log_path, "w", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=unified, restval="")
        writer.writeheader()
        for row in migrated:
            writer.writerow(row)
        writer.writerow(record)


# =========================================================
# 8. SENSITIVITY
# =========================================================

def sensitivity_table(S, K, r, sigma, T, q=0.0):
    base = black_scholes(S, K, r, sigma, T, "call", q)
    rows = []

    for shift in (-0.01, -0.005, 0.0, 0.005, 0.01):
        price = black_scholes(S, K, r + shift, sigma, T, "call", q)
        rows.append({
            "variable": "r",
            "shifted_value": r + shift,
            "call_price": price,
            "delta_vs_base": price - base,
        })

    for shift in (-0.02, -0.01, 0.0, 0.01, 0.02):
        price = black_scholes(S, K, r, sigma + shift, T, "call", q)
        rows.append({
            "variable": "sigma",
            "shifted_value": sigma + shift,
            "call_price": price,
            "delta_vs_base": price - base,
        })

    return pd.DataFrame(rows)


# =========================================================
# 9. WORD REPORT
# =========================================================

def _replace_placeholder(paragraph, values):
    for key, value in values.items():
        while key in paragraph.text:
            texts = [run.text for run in paragraph.runs]
            full = "".join(texts)
            start = full.find(key)
            end = start + len(key)

            pos = 0
            first = last = None
            first_local = last_local = None
            for i, text in enumerate(texts):
                r0, r1 = pos, pos + len(text)
                if first is None and r0 <= start < r1:
                    first = i
                    first_local = start - r0
                if r0 < end <= r1:
                    last = i
                    last_local = end - r0
                    break
                pos = r1

            if first is None:
                break
            if last is None:
                last, last_local = first, first_local + len(key)

            if first == last:
                paragraph.runs[first].text = (
                    texts[first][:first_local] + str(value) + texts[first][last_local:]
                )
            else:
                paragraph.runs[first].text = texts[first][:first_local] + str(value)
                for i in range(first + 1, last):
                    paragraph.runs[i].text = ""
                paragraph.runs[last].text = texts[last][last_local:]


def _iter_report_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p
        for p in section.footer.paragraphs:
            yield p


def generate_word_report(template_path, output_path, values):
    if not HAVE_DOCX:
        raise RuntimeError("python-docx is not installed.")
    if not os.path.isfile(template_path):
        raise FileNotFoundError(f"Report template not found: {template_path}")

    doc = Document(template_path)
    for paragraph in _iter_report_paragraphs(doc):
        _replace_placeholder(paragraph, values)

    unresolved = [
        p.text for p in _iter_report_paragraphs(doc)
        if "{{" in p.text or "}}" in p.text
    ]
    if unresolved:
        raise ValueError("Report contains unresolved placeholders.")

    doc.save(output_path)
    return output_path


def unique_path(path):
    if not os.path.exists(path):
        return path
    base, ext = os.path.splitext(path)
    i = 1
    while os.path.exists(f"{base}_{i}{ext}"):
        i += 1
    candidate = f"{base}_{i}{ext}"
    print(f"[WARN] {path} exists; writing to {candidate}.")
    return candidate


def _fmt_date(d):
    return f"{d.day} {d.strftime('%b %Y')}"


def build_report_values(
    run_date, expiry, vol_start_date, returns, mean_return, sigma,
    S, K, r, q, n, days_to_expiry, published_iv, solved_iv,
    market_ltp, results, skew_stats,
):
    T = days_to_expiry / 365
    solved_valid = np.isfinite(solved_iv)

    values = {
        "{{DATA_DATE}}": _fmt_date(run_date),
        "{{EXPIRY_DATE}}": _fmt_date(expiry),
        "{{DATA_START}}": _fmt_date(vol_start_date) if vol_start_date else "N/A",
        "{{DATA_DATE_DAY}}": str(run_date.day),
        "{{TRADING_DAYS}}": str(len(returns)) if returns is not None else "N/A",
        "{{S}}": f"₹{S:,.2f}",
        "{{ATM_STRIKE}}": f"{K:.0f}",
        "{{R_PCT}}": f"{r * 100:.2f}%",
        "{{Q_PCT}}": f"{q * 100:.2f}%",
        "{{SIGMA}}": f"{sigma * 100:.2f}%" if sigma is not None else "N/A",
        "{{PUBLISHED_IV}}": f"{published_iv * 100:.2f}%",
        "{{SOLVED_IV}}": f"{solved_iv * 100:.2f}%" if solved_valid else "N/A",
        "{{T_DAYS}}": f"{days_to_expiry} days",
        "{{T_DAYS_NUM}}": str(days_to_expiry),
        "{{N_STEPS}}": str(n),
        "{{MEAN_RETURN}}": f"{mean_return * 100:.2f}%" if mean_return is not None else "N/A",
        "{{MARKET_LTP}}": f"₹{market_ltp:.2f}",
    }

    if returns is not None:
        values["{{ROLLING_30}}"] = f"{returns.rolling(30).std().iloc[-1] * np.sqrt(252) * 100:.2f}%"
        values["{{ROLLING_90}}"] = f"{returns.rolling(90).std().iloc[-1] * np.sqrt(252) * 100:.2f}%"
    else:
        values["{{ROLLING_30}}"] = values["{{ROLLING_90}}"] = "N/A"

    price_map = {
        "binomial_call_hist_vol": "{{CALL_HIST}}",
        "binomial_put_hist_vol": "{{PUT_HIST}}",
        "binomial_call_published_iv": "{{CALL_PUBLISHED}}",
        "binomial_put_published_iv": "{{PUT_PUBLISHED}}",
        "binomial_call_solved_iv": "{{CALL_SOLVED}}",
        "binomial_put_solved_iv": "{{PUT_SOLVED}}",
        "bs_call_hist_vol": "{{BS_CALL_HIST}}",
        "bs_put_hist_vol": "{{BS_PUT_HIST}}",
        "bs_call_published_iv": "{{BS_CALL_PUBLISHED}}",
        "bs_put_published_iv": "{{BS_PUT_PUBLISHED}}",
        "bs_call_solved_iv": "{{BS_CALL_SOLVED}}",
        "bs_put_solved_iv": "{{BS_PUT_SOLVED}}",
    }
    for key, placeholder in price_map.items():
        values[placeholder] = f"₹{results[key]:.2f}" if key in results else "N/A"

    # Convergence
    steps = [3, 5, 10, 25, 50, 100, 200, 500]
    if sigma is not None:
        for step in steps:
            price = binomial_price(S, K, r, sigma, T, step, "call", q)
            values[f"{{{{CONV_N{step}}}}}"] = f"{price:.2f}"
        bs_ref = black_scholes(S, K, r, sigma, T, "call", q)
        values["{{CONV_DIFF}}"] = f"{abs(float(values['{{CONV_N500}}']) - bs_ref):.2f}"
    else:
        for step in steps:
            values[f"{{{{CONV_N{step}}}}}"] = "N/A"
        values["{{CONV_DIFF}}"] = "N/A"

    # Put-call parity
    rhs = S * np.exp(-q * T) - K * np.exp(-r * T)
    values["{{PARITY_VAL}}"] = f"{rhs:.2f}"

    parity_map = {
        "HIST": ("binomial_call_hist_vol", "binomial_put_hist_vol"),
        "PUBLISHED": ("binomial_call_published_iv", "binomial_put_published_iv"),
        "SOLVED": ("binomial_call_solved_iv", "binomial_put_solved_iv"),
        "BS_HIST": ("bs_call_hist_vol", "bs_put_hist_vol"),
        "BS_PUBLISHED": ("bs_call_published_iv", "bs_put_published_iv"),
        "BS_SOLVED": ("bs_call_solved_iv", "bs_put_solved_iv"),
    }
    for label, (ckey, pkey) in parity_map.items():
        if ckey in results and pkey in results:
            lhs = results[ckey] - results[pkey]
            values[f"{{{{PARITY_LHS_{label}}}}}"] = f"{lhs:.2f}"
            values[f"{{{{PARITY_DIFF_{label}}}}}"] = f"{abs(lhs - rhs):.2f}"
        else:
            values[f"{{{{PARITY_LHS_{label}}}}}"] = "N/A"
            values[f"{{{{PARITY_DIFF_{label}}}}}"] = "N/A"

    # Greeks: Black-Scholes analytical Greeks
    greek_specs = {
        "HIST": sigma,
        "SOLVED": solved_iv if solved_valid else None,
    }
    for label, vol in greek_specs.items():
        if vol is None:
            for suffix in [
                "DELTA_CALL", "DELTA_PUT", "GAMMA", "VEGA",
                "THETA_CALL", "THETA_PUT", "RHO_CALL", "RHO_PUT",
            ]:
                values[f"{{{{{suffix}_{label}}}}}"] = "N/A"
            continue

        g = bs_greeks(S, K, r, vol, T, q)
        values.update({
            f"{{{{DELTA_CALL_{label}}}}}": f"{g['delta_call']:.3f}",
            f"{{{{DELTA_PUT_{label}}}}}": f"{g['delta_put']:.3f}",
            f"{{{{GAMMA_{label}}}}}": f"{g['gamma']:.4f}",
            f"{{{{VEGA_{label}}}}}": f"{g['vega']:.3f}",
            f"{{{{THETA_CALL_{label}}}}}": f"{g['theta_call']:.3f}",
            f"{{{{THETA_PUT_{label}}}}}": f"{g['theta_put']:.3f}",
            f"{{{{RHO_CALL_{label}}}}}": f"{g['rho_call']:.3f}",
            f"{{{{RHO_PUT_{label}}}}}": f"{g['rho_put']:.3f}",
        })

    # Sensitivity narrative/table: historical volatility base case
    if sigma is not None:
        base = black_scholes(S, K, r, sigma, T, "call", q)
        values["{{SENS_Q_BASE}}"] = f"₹{base:.2f}"
        q_deltas = []
        for label, q_value in [("2", 0.02), ("35", 0.035), ("5", 0.05)]:
            price = black_scholes(S, K, r, sigma, T, "call", q_value)
            delta = price - base
            q_deltas.append(delta)
            values[f"{{{{SENS_Q_{label}}}}}"] = f"{price:.2f}"
            values[f"{{{{SENS_Q_{label}_D}}}}"] = f"{delta:+.3f}"

        vol_delta = black_scholes(S, K, r, sigma + 0.01, T, "call", q) - base
        rate_delta = black_scholes(S, K, r + 0.01, sigma, T, "call", q) - base
        values["{{SENS_VOL_PER_PP}}"] = f"{abs(vol_delta):.2f}"
        values["{{SENS_RATE_PER_PP}}"] = f"{abs(rate_delta):.2f}"
        values["{{SENS_VOL_VS_RATE_RATIO}}"] = (
            f"{abs(vol_delta) / abs(rate_delta):.0f}" if abs(rate_delta) > 1e-9 else "N/A"
        )
        values["{{SENS_Q_RANGE_LOW}}"] = f"{min(abs(x) for x in q_deltas):.2f}"
        values["{{SENS_Q_RANGE_HIGH}}"] = f"{max(abs(x) for x in q_deltas):.2f}"
    else:
        for ph in [
            "{{SENS_VOL_PER_PP}}", "{{SENS_RATE_PER_PP}}", "{{SENS_VOL_VS_RATE_RATIO}}",
            "{{SENS_Q_RANGE_LOW}}", "{{SENS_Q_RANGE_HIGH}}", "{{SENS_Q_BASE}}",
        ]:
            values[ph] = "N/A"
        for label in ["2", "35", "5"]:
            values[f"{{{{SENS_Q_{label}}}}}"] = "N/A"
            values[f"{{{{SENS_Q_{label}_D}}}}"] = "N/A"

    # Volatility skew
    def pct(x):
        return f"{x * 100:.1f}%" if x is not None else "N/A"

    values["{{SKEW_CALL_IV_MIN}}"] = pct(skew_stats.get("call_iv_min"))
    values["{{SKEW_CALL_IV_MAX}}"] = pct(skew_stats.get("call_iv_max"))
    values["{{SKEW_PUT_IV_MIN}}"] = pct(skew_stats.get("put_iv_min"))
    values["{{SKEW_PUT_IV_MAX}}"] = pct(skew_stats.get("put_iv_max"))

    gap = skew_stats.get("atm_gap")
    values["{{SKEW_ATM_GAP}}"] = f"{abs(gap) * 100:.1f}" if gap is not None else "N/A"
    values["{{SKEW_ATM_GAP_STRIKE}}"] = (
        f"{skew_stats['atm_gap_strike']:.0f}" if skew_stats.get("atm_gap_strike") is not None else "N/A"
    )
    values["{{SKEW_ATM_GAP_SIDE}}"] = (
        "put" if gap > 0 else "call" if gap is not None else "N/A"
    )

    if skew_stats.get("filtered"):
        values["{{SKEW_LIQUIDITY_NOTE}}"] = "restricted to strikes that clear the OI/Volume liquidity filter"
    elif skew_stats.get("call_filtered") or skew_stats.get("put_filtered"):
        filtered_side = "call" if skew_stats.get("call_filtered") else "put"
        unfiltered_side = "put" if filtered_side == "call" else "call"
        values["{{SKEW_LIQUIDITY_NOTE}}"] = (
            f"restricted to the liquidity filter on the {filtered_side} side; "
            f"the {unfiltered_side} side used all quoted strikes"
        )
    else:
        values["{{SKEW_LIQUIDITY_NOTE}}"] = "based on all quoted strikes; no OI/Volume columns were available"

    # Directional comparisons for report prose
    def direction(model_price, market_price):
        tolerance = max(0.5, 0.02 * market_price)
        diff = model_price - market_price
        if abs(diff) <= tolerance:
            return "roughly matches"
        return "overprices" if diff > 0 else "underprices"

    def comparison(model_price, market_price):
        tolerance = max(0.5, 0.02 * market_price)
        diff = model_price - market_price
        if abs(diff) <= tolerance:
            return "about the same option value as"
        return "a higher option value than" if diff > 0 else "a lower option value than"

    if "bs_call_hist_vol" in results:
        values["{{HIST_VS_MARKET_DIR}}"] = direction(results["bs_call_hist_vol"], market_ltp)
        values["{{HIST_VS_MARKET_COMPARATIVE}}"] = comparison(results["bs_call_hist_vol"], market_ltp)
    else:
        values["{{HIST_VS_MARKET_DIR}}"] = "cannot be compared to"
        values["{{HIST_VS_MARKET_COMPARATIVE}}"] = "no comparable option value to"

    values["{{PUBLISHED_VS_MARKET_DIR}}"] = (
        direction(results["bs_call_published_iv"], market_ltp)
        if "bs_call_published_iv" in results else "cannot be compared to"
    )

    return values


# =========================================================
# 10. CLI + MAIN
# =========================================================

def build_arg_parser():
    parser = argparse.ArgumentParser(description="TCS options pricing and volatility analysis")
    parser.add_argument("--symbol", default="TCS.NS")
    parser.add_argument("--stock-csv", default="tcs_stock_data.csv")
    parser.add_argument("--chain-csv", default="option_chain_data.csv")
    parser.add_argument("--risk-free-rate", type=float, default=0.0692)
    parser.add_argument("--dividend-yield", type=float, default=0.0)
    parser.add_argument("--n-steps", type=int, default=50)
    parser.add_argument("--expiry", default=None, help="YYYY-MM-DD override")
    parser.add_argument("--expiry-weekday", type=int, default=1, help="Mon=0 ... Sun=6; default Tuesday")
    parser.add_argument("--min-oi", type=int, default=100)
    parser.add_argument("--min-volume", type=int, default=1)
    parser.add_argument("--output-dir", default="outputs")
    parser.add_argument("--log-csv", default="run_history.csv")
    parser.add_argument("--run-date", default=None, help="YYYY-MM-DD; defaults to latest stock CSV date")
    parser.add_argument("--report-template", default="TCS_Options_Pricing_Report_template.docx")
    parser.add_argument("--report-docx", default=None)
    parser.add_argument("--skip-yfinance", action="store_true")
    return parser


def main():
    args = build_arg_parser().parse_args()

    if args.n_steps < 1:
        raise ValueError("--n-steps must be at least 1.")
    if args.min_oi < 0 or args.min_volume < 0:
        raise ValueError("--min-oi and --min-volume cannot be negative.")
    if not 0 <= args.expiry_weekday <= 6:
        raise ValueError("--expiry-weekday must be between 0 and 6.")

    os.makedirs(args.output_dir, exist_ok=True)

    # Stock snapshot
    stock_df = pd.read_csv(args.stock_csv)
    required_stock = {"DATE", "CLOSE"}
    missing = required_stock - set(stock_df.columns)
    if missing:
        raise ValueError(f"Stock CSV missing columns: {sorted(missing)}")

    stock_df["CLOSE"] = pd.to_numeric(
        stock_df["CLOSE"].astype(str).str.replace(",", ""), errors="coerce"
    )
    stock_df["DATE"] = pd.to_datetime(stock_df["DATE"], format="%d-%b-%Y", errors="coerce")
    stock_df = stock_df.dropna(subset=["DATE", "CLOSE"]).sort_values("DATE")
    if stock_df.empty:
        raise ValueError("No valid observations in stock CSV.")

    snapshot_date = stock_df["DATE"].iloc[-1].date()
    run_date = date.fromisoformat(args.run_date) if args.run_date else snapshot_date
    if run_date > snapshot_date:
        raise ValueError(f"Run date {run_date} is after latest stock CSV date {snapshot_date}.")

    stock_row = stock_df[stock_df["DATE"].dt.date == run_date]
    if stock_row.empty:
        raise ValueError(f"No stock-price observation found for run date {run_date}.")
    S = float(stock_row.iloc[0]["CLOSE"])

    # Date/expiry
    expiry = resolve_expiry(run_date, args.expiry, args.expiry_weekday)
    days_to_expiry = (expiry - run_date).days
    T = days_to_expiry / 365
    r, q, n = args.risk_free_rate, args.dividend_yield, args.n_steps

    # Date-based output folder; repeated runs on the same date use _1, _2, ...
    base_run_dir = os.path.join(args.output_dir, run_date.isoformat())
    run_dir = base_run_dir
    suffix = 1
    while os.path.exists(run_dir):
        run_dir = f"{base_run_dir}_{suffix}"
        suffix += 1
    os.makedirs(run_dir)

    print(f"Run output folder = {run_dir}")
    print(f"Exercise style = {EXERCISE_STYLE}")
    print(f"Run date       = {run_date}")
    print(f"Expiry         = {expiry} ({days_to_expiry} days)")

    # Historical volatility
    returns = mean_return = sigma = vol_start_date = None
    if not args.skip_yfinance:
        try:
            returns, mean_return, sigma, vol_start_date = get_historical_volatility(
                args.symbol, end_date=run_date
            )
            print(f"Mean return (descriptive) = {mean_return:.4f}")
            print(f"Historical volatility     = {sigma:.4f}")
            print(f"Volatility data start     = {vol_start_date}")
        except Exception as exc:
            print(f"[WARN] Historical volatility unavailable: {exc}")

    print(f"S              = {S:.2f}")

    # Option chain
    chain, liquidity_info = load_option_chain(args.chain_csv)
    if not liquidity_info["call"]:
        print("[WARN] Call OI/Volume columns missing; strike selection uses the IV/LTP fallback.")
    if not liquidity_info["put"]:
        print("[WARN] Put OI/Volume columns missing; put-side skew is not liquidity-filtered.")

    atm_row, atm_liquidity_filtered = select_atm_strike(
        chain, S, liquidity_info["call"], args.min_oi, args.min_volume
    )
    K = float(atm_row["strike"])
    market_ltp = float(atm_row["call_ltp"])
    published_iv = float(atm_row["call_iv"])

    if not np.isfinite(market_ltp) or market_ltp <= 0:
        raise ValueError("Selected market call LTP is invalid.")
    if not np.isfinite(published_iv) or published_iv <= 0:
        raise ValueError("Selected published call IV is invalid.")

    print(f"Selected K     = {K:.2f}")
    print(f"Market call LTP= {market_ltp:.2f}")
    print(f"Published IV   = {published_iv:.4f}")

    solved_iv = implied_volatility(market_ltp, S, K, r, T, "call", q)
    print(f"Solved IV      = {solved_iv:.4f}" if np.isfinite(solved_iv) else "Solved IV      = N/A")

    # Model prices
    results = {}
    for label, vol in (
        ("hist_vol", sigma),
        ("published_iv", published_iv),
        ("solved_iv", solved_iv),
    ):
        if vol is None or not np.isfinite(vol):
            continue
        results[f"binomial_call_{label}"] = binomial_price(S, K, r, vol, T, n, "call", q)
        results[f"binomial_put_{label}"] = binomial_price(S, K, r, vol, T, n, "put", q)
        results[f"bs_call_{label}"] = black_scholes(S, K, r, vol, T, "call", q)
        results[f"bs_put_{label}"] = black_scholes(S, K, r, vol, T, "put", q)

    print("\n--- Model prices ---")
    for key, value in results.items():
        print(f"{key:30s} = {value:.2f}")

    # Greeks
    if sigma is not None:
        print("\n--- Greeks: historical volatility ---")
        for key, value in bs_greeks(S, K, r, sigma, T, q).items():
            print(f"{key:15s} = {value:.4f}")
    if np.isfinite(solved_iv):
        print("\n--- Greeks: solved IV ---")
        for key, value in bs_greeks(S, K, r, solved_iv, T, q).items():
            print(f"{key:15s} = {value:.4f}")

    # Sensitivity CSV
    if sigma is not None:
        sens_path = os.path.join(run_dir, "sensitivity_analysis.csv")
        sensitivity_table(S, K, r, sigma, T, q).to_csv(sens_path, index=False)
        print(f"Sensitivity table written to {sens_path}")

    skew_stats = compute_skew_stats(
        chain, liquidity_info, K, args.min_oi, args.min_volume
    )

    # Run history
    log_record = {
        "run_timestamp": datetime.now().isoformat(timespec="seconds"),
        "script_version": SCRIPT_VERSION,
        "output_folder": os.path.basename(run_dir),
        "run_date": str(run_date),
        "expiry": str(expiry),
        "S": S,
        "K": K,
        "r": r,
        "q": q,
        "T": T,
        "hist_vol": sigma,
        "published_iv": published_iv,
        "solved_iv": solved_iv,
        "market_ltp": market_ltp,
        "atm_liquidity_filtered": atm_liquidity_filtered,
        **results,
    }
    log_path = os.path.join(args.output_dir, args.log_csv)
    log_run(log_path, log_record)
    print(f"Run appended to {log_path}")

    # Word report
    report_values = build_report_values(
        run_date, expiry, vol_start_date, returns, mean_return, sigma,
        S, K, r, q, n, days_to_expiry, published_iv, solved_iv,
        market_ltp, results, skew_stats,
    )

    report_path = os.path.join(
        run_dir,
        args.report_docx or "TCS_Options_Pricing_Report.docx",
    )
    if os.path.isabs(args.report_docx or ""):
        report_path = args.report_docx
    report_path = unique_path(report_path)

    generate_word_report(args.report_template, report_path, report_values)
    print(f"Word report written to {report_path}")

    # Convergence and rolling-volatility charts
    if sigma is not None:
        steps = [3, 5, 10, 25, 50, 100, 200, 500]
        prices = [binomial_price(S, K, r, sigma, T, step, "call", q) for step in steps]
        bs_ref = black_scholes(S, K, r, sigma, T, "call", q)

        plt.figure(figsize=(8, 5))
        plt.plot(steps, prices, marker="o", label="Binomial Price")
        plt.axhline(bs_ref, linestyle="--", label="Black-Scholes Price")
        plt.xscale("log")
        plt.xlabel("Number of Binomial Steps (log scale)")
        plt.ylabel("Call Option Price (₹)")
        plt.title("Binomial Convergence to Black-Scholes")
        plt.legend()
        plt.grid(True, alpha=0.3)
        plt.tight_layout()
        path = os.path.join(run_dir, "convergence_plot.png")
        plt.savefig(path)
        plt.close()
        print(f"Convergence chart written to {path}")

        rolling_30 = returns.rolling(30).std() * np.sqrt(252)
        rolling_90 = returns.rolling(90).std() * np.sqrt(252)
        plt.figure(figsize=(10, 5))
        plt.plot(rolling_30.index, rolling_30, label="30-Day Rolling Vol")
        plt.plot(rolling_90.index, rolling_90, label="90-Day Rolling Vol")
        plt.axhline(sigma, linestyle="--", label="1-Year Historical Vol")
        plt.axhline(published_iv, linestyle=":", label="Published IV")
        plt.xlabel("Date")
        plt.ylabel("Annualised Volatility")
        plt.title("Rolling Volatility of TCS Returns")
        plt.legend()
        plt.grid(True, alpha=0.3)
        plt.tight_layout()
        path = os.path.join(run_dir, "rolling_volatility.png")
        plt.savefig(path)
        plt.close()
        print(f"Rolling volatility chart written to {path}")

    # Volatility curve uses the same trustworthy-strike filter as report stats.
    call_plot = filter_trustworthy_strikes(
        chain, "call_iv", "call_oi", "call_volume", liquidity_info["call"], K,
        args.min_oi, args.min_volume,
    )
    put_plot = filter_trustworthy_strikes(
        chain, "put_iv", "put_oi", "put_volume", liquidity_info["put"], K,
        args.min_oi, args.min_volume,
    )

    plt.figure(figsize=(11, 5))
    if not call_plot.empty:
        plt.plot(call_plot["strike"], call_plot["call_iv"] * 100, marker="o", label="Call IV")
    if not put_plot.empty:
        plt.plot(put_plot["strike"], put_plot["put_iv"] * 100, marker="o", label="Put IV")
    plt.axvline(S, linestyle="--", label=f"Spot (₹{S:.0f})")
    plt.axvline(K, linestyle=":", label=f"Selected near-ATM strike (₹{K:.0f})")
    plt.xlabel("Strike Price (₹)")
    plt.ylabel("Implied Volatility (%)")
    plt.title(f"TCS Implied Volatility Curve — {expiry.isoformat()} Expiry")
    plt.legend()
    plt.grid(True, alpha=0.3)
    plt.tight_layout()
    path = os.path.join(run_dir, "volatility_smile.png")
    plt.savefig(path)
    plt.close()
    print(f"Volatility smile chart written to {path}")


if __name__ == "__main__":
    main()
